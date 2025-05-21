import csv
import datetime
import json
import logging
import os
import time
from datetime import datetime
from sys import platform

from docx.shared import RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from babel.dates import format_datetime
from mistralai import Mistral
from selenium.common import NoSuchElementException
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.firefox.service import Service
from selenium.webdriver.firefox.options import Options
import docx
import configurations
import re

from configurations import redis_server


# Создание класса SkitConnector
class SkitConnector:
    # Создание конструктора для класса SkitConnector, передача всех аттрибутов
    def __init__(self, url="https://helpdesk.admhmao.ru", login="", passwd="", source=""):
        self.url = url
        self.login = login
        self.passwd = passwd
        self.source = source
        self.driver = None
        self.table = None
        self.dict_source = self.read_source()
        self.doc = docx.Document()

    # Создание метода, использующий driver браузера
    def open_driver(self, path_to_geckodriver, path_to_location_of_firefox):
        # Объект, используемый для настройки параметров запуска браузера
        options = Options()
        # Запуск браузера в фоновом режиме
        options.add_argument("--headless")
        # Путь к исполняемому файлу браузера
        options.binary_location = path_to_location_of_firefox
        # Объект, позволяющий использовать драйвер для работы с браузером
        service = Service(path_to_geckodriver)
        # Назначение используемого браузера, передача аттрибута для работы драйвера, и атрибута, который передаёт настройки параметров запуска браузера
        self.driver = webdriver.Firefox(service=service, options=options)

    # Создание метода, закрывающего driver
    def close_driver(self):
        self.driver.quit()
        self.driver = None

    # Создание метода, создающего словарь с данными
    def read_source(self):
        # Создаём пустой словарь
        result = dict()
        # Создаём цикл, который перечисляет названия файлов source из списка в указанной директории
        for file_name in os.listdir(self.source):
            # Переменная, приводящая название файла в название групп (убирает из названия расширение файла)
            result_file_name = file_name.replace(".csv", "")
            # Словари, являющийся значениями словаря result (ключи result - названия групп)
            result[result_file_name] = dict()
            # Открытие csv файла, находящегося в заданном пути и имеющего определённое название, с отсутствием преобразования символов новой строки, использующий кодировку cp1251 и создание переменной csvfile для работы с файлом csv
            with open(os.path.join(self.source, file_name), newline='', encoding='cp1251') as csvfile:
                # Чтение csv файла, указывание разделяющего знака
                reader = csv.reader(csvfile, delimiter=";")
                # Создание цикла, который перечисляет все строки csv файла
                for row in reader:
                    # Назначение ключей и элементов словарей, являющихся значениями словаря result = {"Группа": "статус заявки" : "ссылка", ...}
                    result[result_file_name][row[0]] = row[1]
        # Создание обратного порядка словаря result
        new_order_dict = {key: result[key] for key in reversed(result.keys())}
        return new_order_dict

    # Метод, авторизующийся на сайте SKIT
    def get_authorization(self):
        # Открытие браузера
        self.driver.get(self.url)
        # Поиск элемента, который используется для ввода логина
        self.driver.find_element("id", "login_name").send_keys(self.login)
        # Поиск элемента, который используется для ввода пароля
        self.driver.find_element("id", "login_password").send_keys(self.passwd)
        # Нажатие на кнопку "Войти"
        self.driver.find_element("name", "submit").click()

    def get_tasks_by_url(self, url):
        self.driver.get(url)
        # Переменная, хранящая в себе элемент, найденный на сайте, в котором записана таблица, работающая с заявками
        table_tasks = self.driver.find_element(By.CLASS_NAME, "tab_cadrehov")
        # Переменная, хранящая в себе элемент, найденный на сайте, в котором находится таблица работы с заявками без заголовков
        tbody = table_tasks.find_element(By.TAG_NAME, "tbody")
        # Переменная, хранящая в себе элементы, найденные на сайте, содержащие строки каждой заявки
        list_of_rows = tbody.find_elements(By.TAG_NAME, "tr")
        answer_list = []
        # Цикл, последовательно обрабатывающий каждую строку, содержащую всю информацию о заявке
        for row in list_of_rows:
            # Переменная, хранящая найденный элемент в строке заявки (столбец)
            cols = row.find_elements(By.TAG_NAME, "td")
            # При отсутствии элементов в строке (когда в фильтре нет заявок)
            if len(cols) == 0:
                break
            # id заявки
            id_task = cols[1].text
            id_task_for_link = id_task.replace(" ", "")
            link_task = f"https://helpdesk.admhmao.ru/front/ticket.form.php?id={id_task_for_link}"
            # Заголовок заявки
            header_task = cols[2].text
            # Дата открытия
            open_data = cols[6].text
            # Последний комментарий
            last_comment = cols[16].text
            # Данные из описания заявки
            try:
                description = cols[15].find_element(By.CLASS_NAME, "fup-popup").get_attribute("textContent")
            except NoSuchElementException:
                description = cols[15].text
            # Все необходимые данные заявки
            current_task = (f"ID заявки: {id_task} \n Дата открытия: {open_data}\nТема: {header_task} \n"
                            f"Текст заявки:{description} \n Последний комментарий: {last_comment}")

            model_answer = configurations.redis_server.get(current_task)
            if model_answer:
                logging.info(f"Ответ из кэша Mistral {id_task}")
                answer_list.append(f"*Заявка*: [{id_task}]({link_task})\n"
                                   f"*Дата заявки*: {open_data}\n"
                                   f"*Краткое содержание заявки*:\n{json.loads(model_answer)}\n"
                                   f"*Последний комментарий*:\n{last_comment}\n"
                                   f"\n")
                continue
            # Использование Mistral
            client = Mistral(api_key=configurations.api_key)
            if "Linux" in platform:
                # Чтение файла prompt
                with open("/prompt.txt", "r",
                          encoding="utf-8") as file:
                    prompt = file.read()
            else:
                with open("C:\\Users\\SoldatovOA\\PycharmProjects\\skit_bot\\src\\prompt.txt", "r",
                          encoding="utf-8") as file:
                    prompt = file.read()
            # Формирование полного ответа с данными заявок и кратким описанием (обработанным текстом нейросетью)
            chat_response = client.chat.complete(
                model=configurations.model,
                temperature=0.4,
                messages=[
                    {
                        "role": "user",
                        "content": prompt + current_task
                    },
                ]
            )
            model_answer = json.dumps(chat_response.choices[0].message.content, ensure_ascii=False)

            redis_server.set(current_task, model_answer, ex=60*60*24*7)

            answer_list.append(f"*Заявка*: [{id_task}]({link_task})\n"
                       f"*Дата заявки*: {open_data}\n"
                       f"*Краткое содержание заявки*:\n{model_answer}\n"
                       f"*Последний комментарий*:\n{last_comment}\n"
                               f"\n")
            logging.info(f"Ответ Mistral {id_task}")
            time.sleep(1)
        return answer_list

    # Метод, формирующий ответ с нужными данными по заявке (ответ даётся не сразу списком, а по готовности по каждой группе)
    def get_report(self):
        find_str = "С 1 по "
        # Словари, содержащие в качестве ключей группы, а в качестве значений - словари, содержащие в качестве ключей статус, а в качестве значений - количество заявок
        source_dict = self.dict_source
        # group_1 = "ЦОП. 3 линия"
        # group_2 = "ЦОП. 3 линия СКУД
        # ...
        for group in source_dict.keys():
            answer_list = []
            answer = "*" + group + "*\n"
            # type_filter_1 = "В работе"
            # type_filter_2 = "Уточнение у инициатора"
            # ...
            for type_filter in source_dict[group].keys():
                self.driver.get(source_dict[group][type_filter])
                try:
                    # Нахождение элемента, переходя по ссылке фильтра (нахождение количества заявок)
                    element = self.driver.find_element(By.XPATH, f"//*[contains(text(), '{find_str}')]")
                    # el_text_i = c 1 по n
                    el_text = element.text
                    # Содержание количества заявок в целом числе count_i = n
                    count = int(el_text.split(" ")[-1])
                except NoSuchElementException:
                    # Если заявок нет (на ссылке фильтра нельзя найти элемент, содержащий количество заявок), присваивается 0
                    count = 0
                # link_1 = ссылка на фильтр "ЦОП. 3 линия, в работе"
                # link_2 = ссылка на фильтр "ЦОП. 3 линия, уточнение у инициатора"
                # ...
                link = source_dict[group][type_filter]
                # answer = группа + статус заявки (гиперссылка в формате Markdown) + количество заявок
                answer += f"[{type_filter}]({link}) {count} \n"
                # Если есть просроченные по ГК, то текст заявки с необходимыми данными выйдет отдельно, включая краткое содержание заявки
                if type_filter == "Просроченные по ГК" and count:
                    model_answer = self.get_tasks_by_url(link)
                    answer_list.extend(model_answer)
            answer_list.insert(0, answer)
            yield answer_list

    def get_report_docx(self):
        logging.info(f"Начата работа report_docx")
        description = {}
        find_str = "С 1 по "
        answer = dict()
        # Словарь типа {"Группа_1": {"Статус_1": "Ссылка_1", "Статус_2": "Ссылка_2", ..., "Статус_n": "Ссылка_n"},
        # "Группа_2": {"Статус_1": "Ссылка_1", "Статус_2": "Ссылка_2", ..., "Статус_n": "Ссылка_n"}, ...,
        # "Группа_n": {"Статус_1": "Ссылка_1", "Статус_2": "Ссылка_2", ..., "Статус_n": "Ссылка_n"}}
        source_dict = self.dict_source
        for group in source_dict.keys():
            logging.info(f"Начата обработка группы {group}")
            current_description = []
            answer[group] = dict()
            # answer = ЦОП. 3 линия ДС
            for type_filter in source_dict[group].keys():
                logging.info(f"Начата обработка статуса заявки {type_filter}")
                self.driver.get(source_dict[group][type_filter])
                try:
                    # element - созданная переменная, содержащая определённый найденный элемент на странице ссылки
                    element = self.driver.find_element(By.XPATH, f"//*[contains(text(), '{find_str}')]")
                    el_text = element.text
                    # Переменная, содержащая количество заявок в целом числе
                    count = int(el_text.split(" ")[-1])
                # Если заявок нет (на ссылке фильтра нельзя найти элемент, содержащий количество заявок), присваивается 0
                except NoSuchElementException:
                    count = 0
                # answer = {Группа_1: {статус_1: кол-во заявок, статус_2: кол-во заявок, ..., статус_n: кол-во заявок},
                # Группа_2: {статус_1: кол-во заявок, статус_2: кол-во заявок, ..., статус_n: кол-во заявок},
                # ... : {...}, Группа_n: {статус_1: кол-во заявок, статус_2: кол-во заявок, ..., статус_n: кол-во заявок}}.
                answer[group][type_filter] = count
                # Ссылка на фильтр
                link = source_dict[group][type_filter]
                if type_filter == "Просроченные по ГК" and count:
                    model_answer = self.get_tasks_by_url(link)
                    current_description.extend(model_answer)
            if len(current_description):
                # description = {Группа_1: id заявки, краткое описание и тд., Группа_2: id заявки, краткое описание и тд., ..., Группа_n: id заявки, краткое описание и тд.}
                description[group] = current_description
        return answer, description

    # Метод, получающий текущую дату и записывающий её в ворд документе
    def data_time(self):
        date_and_time = datetime.now()
        formatted_date_time = format_datetime(date_and_time, format="dd MMMM yyyy г.", locale="ru")
        title_text = self.doc.add_heading(text=f"Количество нерешенных заявок в ГИС Образование на {formatted_date_time}\n", level=1)
        for run in title_text.runs:
            run.font.name = "Times New Roman"
            pt = 13000
            run.font.size = 14 * pt
            run.font.color.rgb = RGBColor(0, 0, 0)
        title_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Метод, обрабатывающий стиль текста названия группы, в которой есть заявки со статусом "просроченные по ГК"
    @staticmethod
    def add_custom_heading(doc, text, level=3, font_name="Times New Roman", font_size=12):
        logging.info(f"Вывод названия группы {text} в файл")
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            pt = 13000
            run.font.name = font_name
            run.font.size = pt*font_size
            run.font.color.rgb = RGBColor(0, 0, 0)

    # Метод, обрабатывающий стиль текста, хранящего информацию по заявкам со статусом "просроченные по ГК"
    @staticmethod
    def add_custom_text(doc, list_text, font_name="Times New Roman", font_size=10):
        logging.info(f"Вывод текста заявки в файл")
        for text in list_text:
            logging.debug(f"Тело заявки {text}")
            text = text.replace("*", "").replace("[", "").replace("]", " ").replace("'", "")
            paragraphs = re.split(r'\\n+', text)
            for paragraph in paragraphs:
                paragraph_text = doc.add_paragraph(paragraph.strip())
                for run in paragraph_text.runs:
                    pt = 13000
                    run.font.name = font_name
                    run.font.size = pt * font_size
                    run.font.color.rgb = RGBColor(0, 0, 0)

    # Метод создания таблицы с количеством заявок по каждой группе и статусам
    def writing_docx(self, name_docx):
        self.doc = docx.Document()

        # Создадим временные метки в файле docx
        self.data_time()
        logging.info(f"Вывод таблицы с количеством заявок в файл")

        # Альбомная ориентация
        section = self.doc.sections[0]
        section.orientation = 1
        section.page_width = docx.shared.Inches(11)  # Ширина в дюймах
        section.page_height = docx.shared.Inches(8.5)  # Высота в дюймах
        dict_answer, description = self.get_report_docx()
        # Создание переменной, содержащей число строк
        count_rows = len(dict_answer)+1
        count_column = 0
        for z in dict_answer.keys():
        #  Создание переменной, содержащей число столбцов
            col = dict_answer[z]
            count_column = len(col)+1
            break
        # Создание таблицы, определение её стиля
        self.table = self.doc.add_table(rows=count_rows, cols=count_column)
        self.table.style = "Table Grid"
        # Получение списка объектов строк, содержащий все строки таблицы (начиная с нулевой строки), содержащие ячейки (объекты "cells"), которые уже могут хранить текст
        rows_table = self.table.rows
        #Текст нулевой ячейки нулевой строки
        rows_table[0].cells[0].text = "Название группы"
        # Введение счётчика строк
        n_row = 1
        name_columns = ["В работе", "Уточнение у инициатора", "Просроченные СКИТ", "Просроченные по ГК", "В ожидании согласования", "Поступило за год", "Поступило за неделю", "Нерешённые"]
        for k in range(0, len(name_columns)):
            rows_table[0].cells[k+1].text = name_columns[k]
        # Цикл, перебирающий ключи dict_answer (Группа_1, Группа_2, ..., Группа_n)
        for n in dict_answer.keys():
            # Получение n-го словаря вида {'статус_n': 'количество'}
            dict_status_and_count = dict_answer[n]
            # Переменная, содержащая объект Row, начиная с первого (смотреть переменную n_row до списка)
            row = rows_table[n_row]
            for g in range(0, len(name_columns)):
                row.cells[g + 1].text = str(dict_status_and_count[name_columns[g]])
            # Создание цикла, проходящего все объекты Cell в n-ом объекте Row
            row.cells[0].text = n
                # Запись в k+1 объект Cell текста, содержащего значение словаря dict_status_and_count (количество заявок, для которых ключи - статусы заявок), обращаясь к словарю через k-ый элемент списка
            n_row += 1
        # Настройка стиля для текста заявок, просроченных по ГК
        for key in description.keys():
            self.add_custom_heading(self.doc, key)
            self.add_custom_text(self.doc, description[key])
        self.doc.save(name_docx)